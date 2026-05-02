"""
增强报告生成器
- 把 Plotly 图表渲染成 PNG(嵌入 HTML/PDF/DOCX)
- 调用 LLM 生成基于所有数据的详细分析
- 支持 Markdown / HTML / PDF / DOCX 四种格式
"""
import base64
import io
import json
from typing import Dict, List, Optional, Tuple
import pandas as pd

from utils.charts import (
    plot_trend, plot_dupont_waterfall, plot_dupont_decomposition,
    plot_peer_comparison, plot_radar
)


# ============================================================
# 图表 → PNG (base64)
# ============================================================
def fig_to_png_base64(fig, width: int = 900, height: int = 450) -> Optional[str]:
    """把 Plotly 图表转换成 base64 PNG 字符串。失败返回 None。"""
    try:
        png_bytes = fig.to_image(format="png", width=width, height=height, scale=2)
        return base64.b64encode(png_bytes).decode("utf-8")
    except Exception:
        # kaleido 没装 / 转换失败时返回 None,后续逻辑跳过
        return None


def fig_to_png_bytes(fig, width: int = 900, height: int = 450) -> Optional[bytes]:
    """图表转 PNG bytes(用于 docx 嵌入)"""
    try:
        return fig.to_image(format="png", width=width, height=height, scale=2)
    except Exception:
        return None


# ============================================================
# 收集所有图表
# ============================================================
def collect_all_charts(ratios: Dict, dupont: Dict, trend_df: pd.DataFrame,
                        compare_df: pd.DataFrame, year: int) -> List[Dict]:
    """
    生成所有图表对象,返回:
    [{"id": "trend_profit", "title": "...", "category": "...",
      "section_anchor": "盈利能力", "fig": <plotly>}, ...]

    section_anchor 用于把图表嵌入到 LLM 报告对应章节
    匹配规则:LLM 报告的二级标题包含 anchor 字符串就把图插到那个章节
    """
    charts = []

    # 1. 杜邦瀑布图 → 嵌入"杜邦"章节
    charts.append({
        "id": "dupont_waterfall",
        "title": f"{year} 年杜邦三因素分解",
        "category": "杜邦分析",
        "section_anchor": "杜邦",
        "fig": plot_dupont_waterfall(dupont, year),
    })

    # 2. 杜邦多年趋势 → 嵌入"杜邦"章节
    if not trend_df.empty:
        dupont_history = {}
        for col in trend_df.columns:
            yr = trend_df[col].to_dict()
            dupont_history[col] = {
                "净利率": yr.get("净利率 (Net Margin)"),
                "总资产周转率": yr.get("总资产周转率"),
                "权益乘数": yr.get("权益乘数 (Equity Multiplier)"),
                "ROE (杜邦计算)": yr.get("ROE 净资产收益率"),
            }
        charts.append({
            "id": "dupont_trend",
            "title": "杜邦三因素历年趋势",
            "category": "杜邦分析",
            "section_anchor": "杜邦",
            "fig": plot_dupont_decomposition(dupont_history),
        })

    # 3. 盈利能力趋势 → 嵌入"盈利"章节
    if not trend_df.empty:
        charts.append({
            "id": "trend_profit",
            "title": "盈利能力趋势(毛利率 / 营业利润率 / 净利率)",
            "category": "趋势分析",
            "section_anchor": "盈利",
            "fig": plot_trend(trend_df,
                ["毛利率 (Gross Margin)", "营业利润率 (Operating Margin)",
                 "净利率 (Net Margin)"], title="盈利能力趋势"),
        })
        # 4. 股东回报趋势 → 嵌入"趋势"或"杜邦"章节
        charts.append({
            "id": "trend_return",
            "title": "股东回报趋势(ROE / ROA)",
            "category": "趋势分析",
            "section_anchor": "趋势",
            "fig": plot_trend(trend_df,
                ["ROE 净资产收益率", "ROA 总资产收益率"], title="股东回报趋势"),
        })
        # 5. 偿债能力趋势 → 嵌入"偿债"章节
        charts.append({
            "id": "trend_solvency",
            "title": "偿债能力趋势(流动 / 速动 / 资产负债率)",
            "category": "趋势分析",
            "section_anchor": "偿债",
            "fig": plot_trend(trend_df,
                ["流动比率", "速动比率", "资产负债率"], title="偿债能力趋势"),
        })

    # 6-9. 同行对比柱状图 + 雷达 → 全部嵌入"同行"章节
    if not compare_df.empty and len(compare_df.columns) >= 2:
        for m in ["净利率 (Net Margin)", "ROE 净资产收益率",
                  "ROA 总资产收益率", "资产负债率"]:
            if m in compare_df.index:
                charts.append({
                    "id": f"peer_{m}",
                    "title": f"同行对比: {m}",
                    "category": "同行对比",
                    "section_anchor": "同行",
                    "fig": plot_peer_comparison(compare_df, m),
                })

        radar_metrics = ["净利率 (Net Margin)", "ROE 净资产收益率",
                         "ROA 总资产收益率", "总资产周转率",
                         "毛利率 (Gross Margin)", "流动比率"]
        charts.append({
            "id": "radar",
            "title": "综合能力雷达图(目标公司 vs 同行均值)",
            "category": "同行对比",
            "section_anchor": "同行",
            "fig": plot_radar(compare_df, radar_metrics),
        })

    return charts


# ============================================================
# 调用 LLM 生成分章节深度分析
# ============================================================
def generate_section_analyses(company_info: Dict, ratios: Dict, dupont: Dict,
                                trend_df: pd.DataFrame, compare_df: pd.DataFrame,
                                year: int, api_key: str,
                                model: str = "gpt-4o-mini") -> Dict[str, str]:
    """
    调用 LLM 为每个章节单独生成基于数据的分析。
    返回:
    {
      "overview": "...",
      "profitability": "...",
      "operating": "...",
      "solvency": "...",
      "dupont": "...",
      "trend": "...",
      "peer": "...",
      "diagnosis": "..."
    }
    """
    from openai import OpenAI
    import json as _json

    # 准备数据
    def clean_ratios(d):
        return {k: round(v, 4) if v is not None else None
                for k, v in d.items() if not k.startswith("_")}

    data = {
        "公司信息": {
            "名称": company_info.get("name"),
            "代码": company_info.get("ticker"),
            "行业": company_info.get("sector"),
            "子行业": company_info.get("industry"),
            "国家": company_info.get("country"),
            "市值(原币)": company_info.get("market_cap"),
            "财报币种": company_info.get("currency"),
            "业务简介": (company_info.get("summary") or "")[:300],
        },
        "财年": year,
        "全部比率": clean_ratios(ratios),
        "杜邦分解": {k: round(v, 4) if v else None for k, v in dupont.items()},
    }

    if not trend_df.empty:
        trend_dict = {}
        for idx in trend_df.index:
            trend_dict[idx] = {str(c): round(v, 4) if pd.notna(v) else None
                                for c, v in trend_df.loc[idx].items()}
        data["历年趋势"] = trend_dict

    if not compare_df.empty:
        peer_dict = {}
        for idx in compare_df.index:
            peer_dict[idx] = {str(c): round(v, 4) if pd.notna(v) else None
                                for c, v in compare_df.loc[idx].items()}
        data["同行对比"] = peer_dict

    system_prompt = """你是一名资深财务分析师。我会给你一家上市公司的完整财务数据,你需要为一份分章节的财务分析报告撰写**8 个章节**的文字分析。

# 关键要求(必读)
1. 用户阅读报告时是这样的体验:**先看图表/数据表,再看你的文字分析**。所以你的文字应该频繁引用具体数字,**像在解读上方的图表/表格**。
2. 不要在文字里说"如下表"、"如下图所示"——因为图表是在你的文字**之前**出现的,要用"上表显示"、"上图可见"、"从图中可以看出"。
3. 绝对不要在文字里贴 Markdown 表格(系统已经渲染了实际表格)。
4. 每段都基于具体数字给出**专业洞察**,而不是简单复述数字。
5. 注意行业语境(银行业杠杆高正常、科技公司毛利率高合理、零售业周转率才是核心等)。
6. 历年趋势数据存在时,务必结合趋势讲"方向"。
7. 同行数据存在时,务必给出"相对位置"判断。

# 8 个章节的分工
- **overview**:公司概览。基于公司信息表和市值,简评公司的市场地位、行业属性、规模量级。约 150-250 字。
- **profitability**:盈利能力分析。基于盈利能力图(毛利率/营业利润率/净利率趋势)和比率表中的盈利相关指标,深入分析盈利质量、利润结构和趋势变化。约 250-350 字。
- **operating**:运营效率。基于资产周转率、存货周转率、应收账款周转率,分析运营效率水平,结合行业特性。约 200-300 字。
- **solvency**:偿债能力与财务结构。基于偿债趋势图(流动比率/速动比率/资产负债率)和指标表,分析短期流动性和长期偿债能力。约 200-300 字。
- **dupont**:杜邦分析。基于杜邦瀑布图(单年三因素)和杜邦历年趋势图(三因素如何变化),拆解 ROE 的核心驱动力,识别哪个因素在变化。约 250-350 字。
- **trend**:历年综合趋势。基于股东回报趋势图(ROE/ROA),结合所有趋势数据,讲发展轨迹和方向性变化。约 200-300 字。
- **peer**:同行业对标。基于同行对比柱状图(净利率/ROE/ROA/资产负债率)和雷达图,分析公司在同行中的相对位置,优劣势。如果没有同行数据,写"未提供同行数据"即可。约 250-350 字。
- **diagnosis**:综合诊断与投资视角。结合所有上述分析,以"### 优势"、"### 风险与不足"、"### 投资者关注要点"三个三级标题给出综合判断。约 300-400 字。

# 输出格式(必须严格遵守)
直接返回 JSON 对象,**不要**用 ```json 代码块包裹,不要任何前后说明文字。结构:
```
{"overview": "文字", "profitability": "文字", "operating": "文字", "solvency": "文字", "dupont": "文字", "trend": "文字", "peer": "文字", "diagnosis": "文字"}
```

每个值是 Markdown 文本(可以用 **粗体**、- 列表、### 三级标题),但不要用 # 一级或 ## 二级标题(系统已经处理章节标题了)。"""

    user_prompt = f"""基于以下数据,为 **{data['公司信息']['名称']}** ({data['公司信息']['代码']}) {year} 年度生成分章节的财务分析。

```json
{_json.dumps(data, ensure_ascii=False, indent=2)}
```

请直接输出 JSON 对象,以 {{ 开始 }} 结束。"""

    client = OpenAI(api_key=api_key)
    resp = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.4,
        max_tokens=4000,
        response_format={"type": "json_object"},  # 强制 JSON
    )
    content = resp.choices[0].message.content or "{}"

    try:
        sections = _json.loads(content)
    except _json.JSONDecodeError:
        # 兜底:返回空 dict,让前端有 fallback
        sections = {}

    # 确保 8 个 key 都有
    expected_keys = ["overview", "profitability", "operating", "solvency",
                     "dupont", "trend", "peer", "diagnosis"]
    for k in expected_keys:
        if k not in sections:
            sections[k] = ""

    return sections


# 兼容旧 API 的封装(把分章节合成一段长文,供下载用)
def generate_llm_analysis(company_info: Dict, ratios: Dict, dupont: Dict,
                            trend_df: pd.DataFrame, compare_df: pd.DataFrame,
                            year: int, api_key: str,
                            model: str = "gpt-4o-mini") -> str:
    """旧 API 兼容封装:返回合并的 Markdown 文本"""
    sections = generate_section_analyses(
        company_info, ratios, dupont, trend_df, compare_df, year, api_key, model
    )
    titles = {
        "overview": "## 一、公司概览",
        "profitability": "## 二、盈利能力分析",
        "operating": "## 三、运营效率",
        "solvency": "## 四、偿债能力与财务结构",
        "dupont": "## 五、杜邦分析:ROE 驱动力拆解",
        "trend": "## 六、历年趋势与发展轨迹",
        "peer": "## 七、同行业对标分析",
        "diagnosis": "## 八、综合诊断与投资视角",
    }
    parts = []
    for key, title in titles.items():
        body = sections.get(key, "").strip()
        if body:
            parts.append(f"{title}\n\n{body}")
    return "\n\n".join(parts)


# ============================================================
# HTML 报告
# ============================================================
HTML_CSS = """
<style>
  body {
    font-family: 'PingFang SC', 'Microsoft YaHei', 'Helvetica Neue', Arial, sans-serif;
    max-width: 920px;
    margin: 32px auto;
    padding: 24px;
    color: #222;
    line-height: 1.7;
    background: #fff;
  }
  h1 { color: #1f4e79; border-bottom: 3px solid #c00000; padding-bottom: 12px; }
  h2 { color: #1f4e79; border-left: 4px solid #c00000; padding-left: 12px;
        margin-top: 36px; }
  h3 { color: #2c5282; margin-top: 24px; }
  .meta { color: #555; font-size: 14px; margin-bottom: 24px; }
  .chart-block { margin: 24px 0; text-align: center; }
  .chart-block img { max-width: 100%; height: auto;
                       border: 1px solid #e5e7eb; border-radius: 6px; }
  .chart-caption { font-size: 13px; color: #666; margin-top: 6px; font-style: italic; }
  table { width: 100%; border-collapse: collapse; margin: 16px 0; font-size: 14px; }
  th, td { border: 1px solid #ddd; padding: 8px 12px; text-align: left; }
  th { background: #f5f7fa; color: #1f4e79; }
  tr:nth-child(even) { background: #fafbfc; }
  .summary-box { background: #f5f7fa; border-left: 4px solid #1f4e79;
                  padding: 16px 20px; margin: 20px 0; border-radius: 4px; }
  .footer { margin-top: 48px; padding-top: 16px; border-top: 1px solid #eee;
              color: #888; font-size: 12px; text-align: center; }
  ul { padding-left: 24px; }
  li { margin: 6px 0; }
  strong { color: #1f4e79; }
</style>
"""


def md_to_html(md: str) -> str:
    """轻量 Markdown 转 HTML(只处理基本元素)"""
    import re
    lines = md.split("\n")
    out = []
    in_list = False
    for line in lines:
        line = line.rstrip()
        if not line:
            if in_list:
                out.append("</ul>")
                in_list = False
            out.append("")
            continue
        # 标题
        m = re.match(r"^(#{1,4})\s+(.+)", line)
        if m:
            if in_list:
                out.append("</ul>")
                in_list = False
            level = len(m.group(1))
            out.append(f"<h{level}>{m.group(2)}</h{level}>")
            continue
        # 列表
        m_li = re.match(r"^[-*]\s+(.+)", line)
        if m_li:
            if not in_list:
                out.append("<ul>")
                in_list = True
            out.append(f"<li>{_inline_md(m_li.group(1))}</li>")
            continue
        # 普通段落
        if in_list:
            out.append("</ul>")
            in_list = False
        out.append(f"<p>{_inline_md(line)}</p>")
    if in_list:
        out.append("</ul>")
    return "\n".join(out)


def _inline_md(text: str) -> str:
    """处理 **bold** *italic* `code`"""
    import re
    text = re.sub(r"\*\*([^\*]+)\*\*", r"<strong>\1</strong>", text)
    text = re.sub(r"\*([^\*]+)\*", r"<em>\1</em>", text)
    text = re.sub(r"`([^`]+)`", r"<code>\1</code>", text)
    return text


def _build_ratios_subset_table(ratios: Dict, keys: List[str]) -> str:
    """生成指定指标子集的 HTML 表格"""
    rows_html = ""
    has_data = False
    for k in keys:
        v = ratios.get(k)
        if v is None and k not in ratios:
            continue
        is_pct = (any(kw in k for kw in ["ROE", "ROA", "ROIC", "Margin"]) or
                  ("率" in k and not any(x in k for x in ["流动比率", "速动比率", "周转率"])))
        if v is None:
            disp = "—"
        elif is_pct:
            disp = f"{v*100:.2f}%"
        elif abs(v) > 1e6:
            disp = f"{v/1e6:.1f}M"
        else:
            disp = f"{v:.3f}"
        rows_html += f"<tr><td>{k}</td><td><strong>{disp}</strong></td></tr>"
        has_data = True
    if not has_data:
        return ""
    return f"""<table><thead><tr><th>指标</th><th>数值</th></tr></thead>
<tbody>{rows_html}</tbody></table>"""


def _build_company_info_table(info: Dict) -> str:
    """公司基本信息表"""
    mcap = info.get("market_cap")
    mcap_str = f"{mcap/1e9:.1f}B {info.get('currency', 'USD')}" if mcap else "—"
    rows = [
        ("公司名称", info.get("name", "—")),
        ("股票代码", info.get("ticker", "—")),
        ("行业", info.get("sector", "—")),
        ("子行业", info.get("industry", "—")),
        ("国家/地区", info.get("country", "—")),
        ("市值", mcap_str),
        ("财报币种", info.get("currency", "—")),
    ]
    rows_html = "".join(f"<tr><td>{k}</td><td>{v}</td></tr>" for k, v in rows)
    return f"<table><tbody>{rows_html}</tbody></table>"


def _build_peer_table(compare_df: pd.DataFrame) -> str:
    """同行对比表 HTML"""
    if compare_df.empty:
        return ""
    html = "<table><thead><tr><th>指标</th>"
    for col in compare_df.columns:
        html += f"<th>{col}</th>"
    html += "</tr></thead><tbody>"
    for idx in compare_df.index:
        is_pct = (any(kw in idx for kw in ["ROE", "ROA", "Margin"]) or
                  ("率" in idx and not any(x in idx for x in ["流动比率", "速动比率", "周转率"])))
        html += f"<tr><td>{idx}</td>"
        for col in compare_df.columns:
            v = compare_df.loc[idx, col]
            if pd.isna(v):
                disp = "—"
            elif is_pct:
                disp = f"{v*100:.2f}%"
            else:
                disp = f"{v:.3f}"
            html += f"<td>{disp}</td>"
        html += "</tr>"
    html += "</tbody></table>"
    return html


def build_html_report(company_info: Dict, year: int, sections: Dict[str, str],
                       charts: List[Dict], ratios: Dict,
                       compare_df: pd.DataFrame) -> str:
    """
    构建 HTML 报告 - 按"图/表 → 文字"模式渲染各章节
    sections: 来自 generate_section_analyses 的 8 个章节文字
    """
    name = company_info.get("name", "Company")
    ticker = company_info.get("ticker", "")
    sector = company_info.get("sector", "—")
    industry = company_info.get("industry", "—")
    currency = company_info.get("currency", "USD")
    mcap = company_info.get("market_cap")
    mcap_str = f"{mcap/1e9:.1f}B {currency}" if mcap else "—"

    # 渲染所有图表为 base64,按 id 索引
    chart_blocks = {}
    for chart in charts:
        b64 = fig_to_png_base64(chart["fig"])
        if b64:
            chart_blocks[chart["id"]] = f"""
<div class="chart-block">
    <img src="data:image/png;base64,{b64}" alt="{chart['title']}" />
    <div class="chart-caption">图: {chart['title']}</div>
</div>"""
        else:
            chart_blocks[chart["id"]] = (
                f'<p class="chart-caption">⚠️ 图表「{chart["title"]}」无法渲染</p>'
            )

    def render_md(text: str) -> str:
        """章节文字 (md) → HTML"""
        if not text or not text.strip():
            return "<p><em>(本章节暂无分析)</em></p>"
        return f'<div class="section-text">{md_to_html(text)}</div>'

    # 各章节使用的指标分组
    profit_keys = ["毛利率 (Gross Margin)", "营业利润率 (Operating Margin)",
                   "净利率 (Net Margin)", "ROA 总资产收益率", "ROE 净资产收益率"]
    operating_keys = ["总资产周转率", "存货周转率", "应收账款周转率"]
    solvency_keys = ["流动比率", "速动比率", "资产负债率", "权益乘数 (Equity Multiplier)",
                     "利息保障倍数"]

    # 找指定 anchor 的图表 id 列表
    def chart_ids_for(anchor: str) -> List[str]:
        return [c["id"] for c in charts if c.get("section_anchor") == anchor]

    def insert_charts(ids: List[str]) -> str:
        return "".join(chart_blocks.get(i, "") for i in ids)

    # ========== 拼装各章节 ==========
    body_html = ""

    # 一、公司概览
    body_html += "<h2>一、公司概览与行业地位</h2>\n"
    body_html += _build_company_info_table(company_info)
    body_html += render_md(sections.get("overview", ""))

    # 二、盈利能力
    body_html += "<h2>二、盈利能力分析</h2>\n"
    body_html += insert_charts(chart_ids_for("盈利"))
    body_html += "<h3>核心盈利指标</h3>"
    body_html += _build_ratios_subset_table(ratios, profit_keys)
    body_html += render_md(sections.get("profitability", ""))

    # 三、运营效率
    body_html += "<h2>三、运营效率</h2>\n"
    body_html += "<h3>运营效率指标</h3>"
    body_html += _build_ratios_subset_table(ratios, operating_keys)
    body_html += render_md(sections.get("operating", ""))

    # 四、偿债能力
    body_html += "<h2>四、偿债能力与财务结构</h2>\n"
    body_html += insert_charts(chart_ids_for("偿债"))
    body_html += "<h3>偿债能力指标</h3>"
    body_html += _build_ratios_subset_table(ratios, solvency_keys)
    body_html += render_md(sections.get("solvency", ""))

    # 五、杜邦分析
    body_html += "<h2>五、杜邦分析:ROE 驱动力拆解</h2>\n"
    body_html += insert_charts(chart_ids_for("杜邦"))
    body_html += render_md(sections.get("dupont", ""))

    # 六、历年趋势
    body_html += "<h2>六、历年趋势与发展轨迹</h2>\n"
    body_html += insert_charts(chart_ids_for("趋势"))
    body_html += render_md(sections.get("trend", ""))

    # 七、同行业对标
    body_html += "<h2>七、同行业对标分析</h2>\n"
    if not compare_df.empty:
        body_html += "<h3>同行对照表</h3>"
        body_html += _build_peer_table(compare_df)
        body_html += insert_charts(chart_ids_for("同行"))
    body_html += render_md(sections.get("peer", ""))

    # 八、综合诊断
    body_html += "<h2>八、综合诊断与投资视角</h2>\n"
    body_html += render_md(sections.get("diagnosis", ""))

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>{name} ({ticker}) {year} 年财务分析报告</title>
{HTML_CSS}
</head>
<body>

<h1>{name} ({ticker})</h1>
<p class="meta">
<strong>{year} 年度财务分析报告</strong><br>
行业: {sector} / {industry} &nbsp;|&nbsp; 国家: {company_info.get('country', '—')} \
&nbsp;|&nbsp; 市值: {mcap_str}
</p>

<div class="summary-box">
本报告由 AI 财务分析 Agent 自动生成。每个章节先呈现关键数据(图表/表格),\
再由 AI 基于这些数据展开专业分析。
</div>

{body_html}

<div class="footer">
本报告基于 Yahoo Finance 公开财务数据自动生成,仅供学术与参考用途,不构成投资建议。<br>
Generated by Financial Analysis AI Agent.
</div>

</body>
</html>"""


# ============================================================
# DOCX 报告
# ============================================================
def build_docx_report(company_info: Dict, year: int, sections: Dict[str, str],
                       charts: List[Dict], ratios: Dict,
                       compare_df: pd.DataFrame) -> bytes:
    """构建 Word 报告 → bytes(按章节+图表先于文字模式)"""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 标题
    doc.add_heading(
        f"{company_info.get('name', '')} ({company_info.get('ticker', '')})", level=0)
    subtitle = doc.add_paragraph(f"{year} 年度财务分析报告")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 元信息
    meta_p = doc.add_paragraph()
    meta_p.add_run(f"行业: {company_info.get('sector', '—')} / "
                    f"{company_info.get('industry', '—')}\n").italic = True
    meta_p.add_run(f"国家: {company_info.get('country', '—')}  |  ").italic = True
    mcap = company_info.get("market_cap")
    mcap_str = f"{mcap/1e9:.1f}B {company_info.get('currency', 'USD')}" if mcap else "—"
    meta_p.add_run(f"市值: {mcap_str}").italic = True

    doc.add_paragraph(
        "本报告由 AI 财务分析 Agent 自动生成。每章节先呈现关键数据(图表/表格),"
        "再由 AI 基于这些数据展开专业分析。"
    )

    # ========== 辅助函数 ==========
    def add_md_paragraphs(text: str):
        """把 Markdown 文字段落写入 docx,处理 **粗体**、列表、### 标题"""
        if not text or not text.strip():
            doc.add_paragraph("(本章节暂无分析)").italic = True
            return
        for line in text.split("\n"):
            line = line.rstrip()
            if not line.strip():
                continue
            if line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("- ") or line.startswith("* "):
                p = doc.add_paragraph(style="List Bullet")
                _add_runs_with_bold(p, line[2:])
            else:
                p = doc.add_paragraph()
                _add_runs_with_bold(p, line)

    def add_ratios_table(keys: List[str]):
        """添加指定指标子集的表格"""
        valid_rows = []
        for k in keys:
            v = ratios.get(k)
            if v is None and k not in ratios:
                continue
            is_pct = (any(kw in k for kw in ["ROE", "ROA", "ROIC", "Margin"]) or
                      ("率" in k and not any(x in k for x in ["流动比率", "速动比率", "周转率"])))
            if v is None:
                disp = "—"
            elif is_pct:
                disp = f"{v*100:.2f}%"
            elif abs(v) > 1e6:
                disp = f"{v/1e6:.1f}M"
            else:
                disp = f"{v:.3f}"
            valid_rows.append((k, disp))
        if not valid_rows:
            return
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "指标"
        hdr[1].text = "数值"
        for k, disp in valid_rows:
            row = table.add_row().cells
            row[0].text = k
            row[1].text = disp

    def add_charts_for(anchor: str):
        """插入指定 anchor 的所有图表"""
        for ch in charts:
            if ch.get("section_anchor") != anchor:
                continue
            png = fig_to_png_bytes(ch["fig"])
            if png:
                doc.add_picture(io.BytesIO(png), width=Inches(6.0))
                cap = doc.add_paragraph(f"图: {ch['title']}")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cap.runs:
                    run.italic = True
                    run.font.size = Pt(9)
            else:
                p = doc.add_paragraph(f"⚠️ 图表「{ch['title']}」无法渲染")
                for run in p.runs:
                    run.italic = True

    def add_company_info_table():
        """公司基本信息表"""
        rows_data = [
            ("公司名称", company_info.get("name", "—")),
            ("股票代码", company_info.get("ticker", "—")),
            ("行业", company_info.get("sector", "—")),
            ("子行业", company_info.get("industry", "—")),
            ("国家/地区", company_info.get("country", "—")),
            ("市值", mcap_str),
            ("财报币种", company_info.get("currency", "—")),
        ]
        table = doc.add_table(rows=len(rows_data), cols=2)
        table.style = "Light Grid Accent 1"
        for i, (k, v) in enumerate(rows_data):
            table.rows[i].cells[0].text = k
            table.rows[i].cells[1].text = str(v)

    def add_peer_table():
        """同行对照表"""
        if compare_df.empty:
            return
        table = doc.add_table(rows=1, cols=len(compare_df.columns) + 1)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "指标"
        for i, col in enumerate(compare_df.columns):
            hdr[i + 1].text = str(col)[:25]
        for idx in compare_df.index:
            is_pct = (any(kw in idx for kw in ["ROE", "ROA", "Margin"]) or
                      ("率" in idx and not any(x in idx for x in ["流动比率", "速动比率", "周转率"])))
            row = table.add_row().cells
            row[0].text = idx
            for i, col in enumerate(compare_df.columns):
                v = compare_df.loc[idx, col]
                if pd.isna(v):
                    row[i + 1].text = "—"
                elif is_pct:
                    row[i + 1].text = f"{v*100:.2f}%"
                else:
                    row[i + 1].text = f"{v:.3f}"

    # 各章节使用的指标分组
    profit_keys = ["毛利率 (Gross Margin)", "营业利润率 (Operating Margin)",
                   "净利率 (Net Margin)", "ROA 总资产收益率", "ROE 净资产收益率"]
    operating_keys = ["总资产周转率", "存货周转率", "应收账款周转率"]
    solvency_keys = ["流动比率", "速动比率", "资产负债率", "权益乘数 (Equity Multiplier)",
                     "利息保障倍数"]

    # ========== 一、公司概览 ==========
    doc.add_heading("一、公司概览与行业地位", level=1)
    add_company_info_table()
    add_md_paragraphs(sections.get("overview", ""))

    # ========== 二、盈利能力 ==========
    doc.add_heading("二、盈利能力分析", level=1)
    add_charts_for("盈利")
    doc.add_heading("核心盈利指标", level=2)
    add_ratios_table(profit_keys)
    add_md_paragraphs(sections.get("profitability", ""))

    # ========== 三、运营效率 ==========
    doc.add_heading("三、运营效率", level=1)
    doc.add_heading("运营效率指标", level=2)
    add_ratios_table(operating_keys)
    add_md_paragraphs(sections.get("operating", ""))

    # ========== 四、偿债能力 ==========
    doc.add_heading("四、偿债能力与财务结构", level=1)
    add_charts_for("偿债")
    doc.add_heading("偿债能力指标", level=2)
    add_ratios_table(solvency_keys)
    add_md_paragraphs(sections.get("solvency", ""))

    # ========== 五、杜邦分析 ==========
    doc.add_heading("五、杜邦分析:ROE 驱动力拆解", level=1)
    add_charts_for("杜邦")
    add_md_paragraphs(sections.get("dupont", ""))

    # ========== 六、历年趋势 ==========
    doc.add_heading("六、历年趋势与发展轨迹", level=1)
    add_charts_for("趋势")
    add_md_paragraphs(sections.get("trend", ""))

    # ========== 七、同行业对标 ==========
    doc.add_heading("七、同行业对标分析", level=1)
    if not compare_df.empty:
        doc.add_heading("同行对照表", level=2)
        add_peer_table()
        add_charts_for("同行")
    add_md_paragraphs(sections.get("peer", ""))

    # ========== 八、综合诊断 ==========
    doc.add_heading("八、综合诊断与投资视角", level=1)
    add_md_paragraphs(sections.get("diagnosis", ""))

    # 保存
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_runs_with_bold(paragraph, text):
    """处理 **粗体** 渲染到 docx"""
    import re
    parts = re.split(r"(\*\*[^\*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


# ============================================================
# PDF 报告 - 通过 HTML 转换
# ============================================================
def build_pdf_report(html_content: str) -> Optional[bytes]:
    """
    HTML → PDF。使用 weasyprint(若可用)或 fpdf 兜底。
    Streamlit Cloud 默认不带 weasyprint 的系统依赖,所以可能失败。
    失败时返回 None,前端应回退提示用户用 HTML/DOCX。
    """
    # 方案 1:weasyprint(最佳质量,需系统库)
    try:
        from weasyprint import HTML
        return HTML(string=html_content).write_pdf()
    except Exception:
        pass

    # 方案 2:无可用工具
    return None
